"""Reports routes — export data as CSV / Excel / Word / PDF / TXT."""
import io
import csv
import datetime as dt
from typing import List, Optional, Tuple

from fastapi import APIRouter, Depends, HTTPException, Query, Request, Response
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from ..database import get_db
from .. import models, auth
from ..utils.audit import record_audit
from ..utils.crypto import decrypt
from ..utils.validation import mask_aadhaar

router = APIRouter(prefix="/reports", tags=["reports"])


# ---------------------------------------------------------------------------
# Supported formats & content-types
# ---------------------------------------------------------------------------
FORMATS = {"csv", "xlsx", "docx", "pdf", "txt"}

MIME = {
    "csv":  "text/csv",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf":  "application/pdf",
    "txt":  "text/plain",
}


# ---------------------------------------------------------------------------
# Data fetchers — one per report type. Returns (headers, rows, title).
# ---------------------------------------------------------------------------
def _residents(db: Session):
    rows = db.query(models.Resident).order_by(models.Resident.created_at.desc()).all()
    headers = ["Resident ID", "First Name", "Last Name", "Gender", "Age",
               "Mobile", "Email", "Village", "Ward", "Category", "Family ID", "Created At"]
    data = [[r.resident_id, r.first_name, r.last_name, r.gender, r.age,
             r.mobile_number, r.email, r.village, r.ward_number, r.category,
             r.family_id, r.created_at.isoformat() if r.created_at else ""] for r in rows]
    return headers, data, "Residents Report"


def _families(db: Session):
    rows = db.query(models.Family).order_by(models.Family.created_at.desc()).all()
    headers = ["Family ID", "Head Name", "Village", "Ward", "Created At"]
    data = [[f.family_id, f.head_name, f.village, f.ward_number,
             f.created_at.isoformat() if f.created_at else ""] for f in rows]
    return headers, data, "Families Report"


def _schemes(db: Session):
    rows = db.query(models.Scheme).order_by(models.Scheme.created_at.desc()).all()
    headers = ["Name", "Status", "Age Min", "Age Max", "Gender", "Category",
               "Income Limit", "Deadline"]
    data = [[s.name, s.status, s.age_min, s.age_max, s.gender, s.category,
             s.income_limit,
             s.application_deadline.isoformat() if s.application_deadline else ""]
            for s in rows]
    return headers, data, "Government Schemes Report"


def _applications(db: Session, status_filter: str = ""):
    query = db.query(models.Application)
    if status_filter:
        query = query.filter(models.Application.status == status_filter)
    rows = query.order_by(models.Application.created_at.desc()).all()
    headers = ["Application Number", "Resident", "Scheme", "Status",
               "Created At", "Updated At"]
    data = []
    for a in rows:
        resident_name = ""
        if a.resident:
            resident_name = f"{a.resident.first_name} {a.resident.last_name or ''}".strip()
        data.append([a.application_number, resident_name,
                     a.scheme.name if a.scheme else "",
                     a.status,
                     a.created_at.isoformat() if a.created_at else "",
                     a.updated_at.isoformat() if a.updated_at else ""])
    title = "Applications Report"
    if status_filter:
        title += f" — {status_filter.title()}"
    return headers, data, title


def _ward(db: Session):
    from sqlalchemy import func
    rows = db.query(
        models.Resident.village,
        models.Resident.ward_number,
        func.count(models.Resident.id),
    ).group_by(models.Resident.village, models.Resident.ward_number).all()
    headers = ["Village", "Ward", "Resident Count"]
    data = [list(r) for r in rows]
    return headers, data, "Ward Report"


def _village(db: Session):
    from sqlalchemy import func
    rows = db.query(
        models.Resident.village,
        func.count(models.Resident.id),
    ).group_by(models.Resident.village).all()
    headers = ["Village", "Resident Count"]
    data = [list(r) for r in rows]
    return headers, data, "Village Report"


FETCHERS = {
    "residents":    _residents,
    "families":     _families,
    "schemes":      _schemes,
    "applications": _applications,
    "ward":         _ward,
    "village":      _village,
}


# ---------------------------------------------------------------------------
# Format renderers — each returns (bytes, content_type, filename_suffix)
# ---------------------------------------------------------------------------
def _to_csv(headers: List[str], rows: List[list], title: str) -> bytes:
    buf = io.StringIO()
    writer = csv.writer(buf)
    writer.writerow([title])
    writer.writerow([f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}"])
    writer.writerow([])
    writer.writerow(headers)
    for r in rows:
        writer.writerow(r)
    return buf.getvalue().encode("utf-8")


def _to_txt(headers: List[str], rows: List[list], title: str) -> bytes:
    lines = [title, "=" * len(title), ""]
    lines.append(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    # Compute column widths
    widths = [len(h) for h in headers]
    for r in rows:
        for i, v in enumerate(r):
            if i < len(widths):
                widths[i] = max(widths[i], len(str(v if v is not None else "")))
    # Header row
    lines.append("  ".join(str(h).ljust(widths[i]) for i, h in enumerate(headers)))
    lines.append("  ".join("-" * w for w in widths))
    # Data rows
    for r in rows:
        lines.append("  ".join(str(v if v is not None else "").ljust(widths[i])
                                for i, v in enumerate(r)))
    return "\n".join(lines).encode("utf-8")


def _to_xlsx(headers: List[str], rows: List[list], title: str) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = title[:31]  # Excel sheet name limit

    # Title row
    ws.cell(row=1, column=1, value=title)
    ws.cell(row=1, column=1).font = Font(size=14, bold=True, color="1F2937")
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))

    # Meta row
    ws.cell(row=2, column=1, value=f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    ws.cell(row=2, column=1).font = Font(size=9, italic=True, color="6B7280")
    ws.merge_cells(start_row=2, start_column=1, end_row=2, end_column=len(headers))

    # Header row
    header_row = 4
    header_fill = PatternFill("solid", fgColor="4F46E5")
    header_font = Font(size=11, bold=True, color="FFFFFF")
    thin = Side(border_style="thin", color="D1D5DB")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    for i, h in enumerate(headers, start=1):
        cell = ws.cell(row=header_row, column=i, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    # Data rows
    for r_idx, r in enumerate(rows, start=header_row + 1):
        for c_idx, v in enumerate(r, start=1):
            cell = ws.cell(row=r_idx, column=c_idx, value=v if v is not None else "")
            cell.border = border
            cell.alignment = Alignment(vertical="top", wrap_text=False)

    # Auto column widths
    for i, h in enumerate(headers, start=1):
        col_letter = get_column_letter(i)
        max_len = len(str(h))
        for r in rows:
            if i - 1 < len(r):
                max_len = max(max_len, len(str(r[i - 1] if r[i - 1] is not None else "")))
        ws.column_dimensions[col_letter].width = min(max_len + 2, 50)

    # Freeze header
    ws.freeze_panes = ws.cell(row=header_row + 1, column=1)

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def _to_docx(headers: List[str], rows: List[list], title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    h = doc.add_heading(title, level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Meta
    meta = doc.add_paragraph()
    meta_run = meta.add_run(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    meta_run.italic = True

    doc.add_paragraph()  # spacer

    # Table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    # Header row
    for i, hdr in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(hdr)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, r in enumerate(rows, start=1):
        for c_idx, v in enumerate(r):
            if c_idx < len(headers):
                table.rows[r_idx].cells[c_idx].text = str(v if v is not None else "")

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _to_pdf(headers: List[str], rows: List[list], title: str) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.enums import TA_LEFT

    buf = io.BytesIO()
    page_size = landscape(A4)
    doc = SimpleDocTemplate(buf, pagesize=page_size,
                            topMargin=15 * mm, bottomMargin=15 * mm,
                            leftMargin=12 * mm, rightMargin=12 * mm)

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("Title", parent=styles["Title"],
                                  fontSize=16, textColor=colors.HexColor("#1F2937"),
                                  spaceAfter=4)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.HexColor("#6B7280"),
                                 spaceAfter=12)
    cell_style = ParagraphStyle("Cell", parent=styles["Normal"],
                                 fontSize=8, leading=10, alignment=TA_LEFT)
    header_style = ParagraphStyle("Hdr", parent=cell_style,
                                   fontSize=9, textColor=colors.white,
                                   fontName="Helvetica-Bold")

    story = [
        Paragraph(title, title_style),
        Paragraph(f"Generated: {dt.datetime.now().strftime('%Y-%m-%d %H:%M')}", meta_style),
    ]

    # Wrap every cell in a Paragraph so long text wraps cleanly
    table_data = [[Paragraph(str(h), header_style) for h in headers]]
    for r in rows:
        table_data.append([Paragraph(str(v if v is not None else ""), cell_style) for v in r])

    # Compute column widths proportionally to page width
    page_width = page_size[0] - 24 * mm
    col_count = max(len(headers), 1)
    col_w = page_width / col_count
    col_widths = [col_w] * col_count

    tbl = Table(table_data, colWidths=col_widths, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4F46E5")),
        ("TEXTCOLOR",  (0, 0), (-1, 0), colors.white),
        ("ALIGN",      (0, 0), (-1, -1), "LEFT"),
        ("VALIGN",     (0, 0), (-1, -1), "TOP"),
        ("GRID",       (0, 0), (-1, -1), 0.25, colors.HexColor("#D1D5DB")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
            [colors.white, colors.HexColor("#F3F4F6")]),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(tbl)

    doc.build(story)
    return buf.getvalue()


RENDERERS = {
    "csv":  _to_csv,
    "txt":  _to_txt,
    "xlsx": _to_xlsx,
    "docx": _to_docx,
    "pdf":  _to_pdf,
}


# ---------------------------------------------------------------------------
# Unified endpoint: GET /reports/{entity}.{format}
# ---------------------------------------------------------------------------
@router.get("/{entity}.{fmt}")
def export_report(
    entity: str,
    fmt: str,
    request: Request,
    status_filter: str = "",
    db: Session = Depends(get_db),
    user: models.User = Depends(auth.require_admin_or_operator),
):
    """Generate a report in the requested format.

    Path params:
      entity: residents | families | schemes | applications | ward | village
      fmt:    csv | xlsx | docx | pdf | txt

    Query params:
      status_filter: only for `applications` (applied/pending/processing/approved/rejected)
    """
    fmt = fmt.lower()
    if fmt not in FORMATS:
        raise HTTPException(status_code=400,
                            detail=f"Unsupported format '{fmt}'. Use one of: {', '.join(sorted(FORMATS))}")
    if entity not in FETCHERS:
        raise HTTPException(status_code=404,
                            detail=f"Unknown report '{entity}'. Available: {', '.join(sorted(FETCHERS.keys()))}")

    fetcher = FETCHERS[entity]
    if entity == "applications":
        headers, rows, title = fetcher(db, status_filter=status_filter)
    else:
        headers, rows, title = fetcher(db)

    renderer = RENDERERS[fmt]
    body = renderer(headers, rows, title)

    filename = f"{entity}_{dt.datetime.now().strftime('%Y%m%d_%H%M')}.{fmt}"
    record_audit(db, user, "export_report", entity, None,
                 f"Exported {entity} report as {fmt.upper()} ({len(rows)} rows)", request)

    return Response(
        content=body,
        media_type=MIME[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Custom list endpoint: POST /reports/residents/custom.{format}
# Body: { "ids": [1, 2, 3], "title": "Optional custom title" }
# Returns a report containing only the specified residents.
# ---------------------------------------------------------------------------
from pydantic import BaseModel as PydanticBaseModel, Field as PydanticField


class CustomListRequest(PydanticBaseModel):
    ids: List[int] = PydanticField(..., min_length=1, max_length=10000)
    title: Optional[str] = None


def _residents_by_ids(db: Session, ids: List[int]):
    """Fetch a specific set of residents by ID, preserving the given order."""
    rows = db.query(models.Resident).filter(models.Resident.id.in_(ids)).all()
    # Preserve the order of the input IDs
    by_id = {r.id: r for r in rows}
    ordered = [by_id[i] for i in ids if i in by_id]

    headers = ["Resident ID", "First Name", "Last Name", "Gender", "Age",
               "Mobile", "Email", "Village", "Ward", "Category", "Family ID", "Created At"]
    data = [[r.resident_id, r.first_name, r.last_name, r.gender, r.age,
             r.mobile_number, r.email, r.village, r.ward_number, r.category,
             r.family_id, r.created_at.isoformat() if r.created_at else ""]
            for r in ordered]
    return headers, data, "Custom Resident List"


@router.post("/residents/custom.{fmt}")
def export_custom_residents(
    fmt: str,
    payload: CustomListRequest,
    request: Request,
    db: Session = Depends(get_db),
    user: models.User = Depends(auth.require_admin_or_operator),
):
    """Generate a report for a custom selection of residents.

    Body:
      { "ids": [1, 2, 3], "title": "Beneficiaries for PM Awas Yojana" }

    Use case: user selects specific residents on the list page and downloads
    only those records — e.g. to share a list of eligible beneficiaries with
    a government office.
    """
    fmt = fmt.lower()
    if fmt not in FORMATS:
        raise HTTPException(status_code=400,
                            detail=f"Unsupported format '{fmt}'. Use one of: {', '.join(sorted(FORMATS))}")

    headers, rows, default_title = _residents_by_ids(db, payload.ids)
    title = payload.title or default_title

    renderer = RENDERERS[fmt]
    body = renderer(headers, rows, title)

    filename = f"custom_residents_{dt.datetime.now().strftime('%Y%m%d_%H%M')}.{fmt}"
    record_audit(db, user, "export_report", "residents", None,
                 f"Exported custom resident list as {fmt.upper()} ({len(rows)} of {len(payload.ids)} requested IDs)", request)

    return Response(
        content=body,
        media_type=MIME[fmt],
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Backward-compatible CSV routes (keep old links working)
# ---------------------------------------------------------------------------
@router.get("/residents.csv")
def residents_csv(db: Session = Depends(get_db),
                  user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _residents(db)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="residents.csv"'})


@router.get("/families.csv")
def families_csv(db: Session = Depends(get_db),
                 user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _families(db)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="families.csv"'})


@router.get("/schemes.csv")
def schemes_csv(db: Session = Depends(get_db),
                user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _schemes(db)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="schemes.csv"'})


@router.get("/applications.csv")
def applications_csv(status_filter: str = "",
                     db: Session = Depends(get_db),
                     user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _applications(db, status_filter=status_filter)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="applications.csv"'})


@router.get("/ward.csv")
def ward_csv(db: Session = Depends(get_db),
             user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _ward(db)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="ward_report.csv"'})


@router.get("/village.csv")
def village_csv(db: Session = Depends(get_db),
                user: models.User = Depends(auth.require_admin_or_operator)):
    headers, rows, title = _village(db)
    return Response(content=_to_csv(headers, rows, title), media_type=MIME["csv"],
                    headers={"Content-Disposition": 'attachment; filename="village_report.csv"'})
